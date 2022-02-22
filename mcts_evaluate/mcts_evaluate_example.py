#!/usr/bin/env python 
# -*- coding: utf-8 -*-
# @Time    : 2021/9/30 10:28
# @Author  : zhangbc0315@outlook.com
# @File    : mcts_evaluate_example.py
# @Software: PyCharm
import os

from docx import Document
from docx.shared import Inches
import pandas as pd
from rdkit.Chem import AllChem
from rdkit.Chem import Draw

from mcts_evaluate.mcts_evaluate_count import MctsEvaluateCount, MctsPath


class MctsEvaluateExample:

    @classmethod
    def _kekulize_rxn(cls, rxn: AllChem.ChemicalReaction):
        for r in rxn.GetReactants():
            AllChem.Kekulize(r, clearAromaticFlags=True)
        for p in rxn.GetProducts():
            AllChem.Kekulize(p, clearAromaticFlags=True)

    @classmethod
    def _init_rxn(cls, rxn: AllChem.ChemicalReaction):
        new_rxn = AllChem.ChemicalReaction()
        for r in rxn.GetReactants():
            new_r = AllChem.MolFromSmiles(AllChem.MolToSmiles(r))
            new_rxn.AddReactantTemplate(new_r)
        for p in rxn.GetProducts():
            new_p = AllChem.MolFromSmiles(AllChem.MolToSmiles(p))
            new_rxn.AddProductTemplate(new_p)
        return new_rxn

    @classmethod
    def display_rxns(cls, rxn_path, doc, ei):
        if pd.isna(rxn_path):
            return False
        rxn_path = eval(rxn_path)
        rxn_path.reverse()
        for j, rxn_str in enumerate(rxn_path):

            rxn_smiles = rxn_str.split(' |c| ')[0]
            # doc.add_paragraph(f'{rxn_smiles}', style='List Number')
            rxn = AllChem.ReactionFromSmarts(rxn_smiles)
            rxn = cls._init_rxn(rxn)
            # cls._kekulize_rxn(rxn)
            pic = Draw.ReactionToImage(rxn)
            pic_fp = os.path.join(MctsPath.ALL_EXAMPLE_DP, f'{ei}-{j}.png')
            pic.save(pic_fp)
            doc.add_picture(pic_fp, width=Inches(6.5))

        return True

    @classmethod
    def process(cls):
        df = MctsEvaluateCount.get_eva_df()
        df = df.reset_index()
        df['rxn_num'] = df['rxn_path'].map(lambda x: 0 if pd.isna(x) else len(eval(x)))
        df = df.sort_values(by='rxn_num', ascending=False)
        df = df.reset_index()
        doc = Document()
        num = 0
        for i, row in df.iterrows():
            doc.add_heading(f'{i} - {row["inchi"]}', level=2)
            res = cls.display_rxns(row['rxn_path'], doc, i)
            if res:
                num += 1
            if num >= 100:
                break
        doc.save(MctsPath.ALL_EXAMPLE_FP)


if __name__ == "__main__":
    MctsEvaluateExample.process()
